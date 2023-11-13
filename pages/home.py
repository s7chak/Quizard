import json
import json
import os
import sys

import dash
import dash_bootstrap_components as dbc
from dash import dcc, callback
from dash import html
from dash.dependencies import Input, Output, State
from docx import Document

from config import Config
from ops.opapp import Util, LanguageModel

dash.register_page(__name__, path='/home')

print("Home page opened.")

links = []
num = '2'
quiz = {}
def serve_layout():
    layout = html.Div([
        dbc.Container([
            html.P(),
            dbc.Col(html.H4(children="Train Quiz Agent", className="mb-2")),
            html.P(),
            html.Div(id='read_status'),
            html.Div(id='process_status'),
            html.Div(id='quiz_status'),
            dbc.Row([
                dbc.Col(
                    html.Div(id='input-container', children=[
                            html.P(
                                dcc.Input(id='input-now', type='text', value='', debounce=True, style= {
                                    'width': '50%',
                                    'height': '20%'
                                })
                            ),
                        ]
                    ),
                ),
                dbc.Col(
                    html.Div(id='input-container-2', style={
                                'alignItems':'center'},
                children=[
                        html.P(
                            dcc.Input(id='input-question-count', placeholder='Total Questions', type='text', value='',
                                      debounce=True, style={
                                    'width': '30%',
                                    'height': '10%'
                                })
                        ),
                        html.P(
                            dcc.Input(id='low-question-count', placeholder='Easy', type='text', value='', debounce=True, style={
                                'width': '30%',
                                'height': '10%',
                            })
                        ),
                        html.P(
                            dcc.Input(id='mid-question-count', placeholder='Medium', type='text', value='', debounce=True, style={
                                'width': '30%',
                                'height': '10%'
                            })
                        ),
                        html.P(
                            dcc.Input(id='high-question-count', placeholder='Hard', type='text', value='', debounce=True, style={
                                'width': '30%',
                                'height': '10%'
                            })
                        ),
                    ]),
                ),
            ]),
            dbc.Col(dbc.Button("Submit", id='submit_button')),
            dbc.Col(dbc.Button("Show Quiz", id='show_button')),
            html.Div([
                dcc.Input(id='quiz_name', placeholder='Name Quiz', type='text', value='', debounce=True, style={
                    'width': '20%',
                    'height': '10%'
                }),
                dcc.Download(id='download_link'),
                dbc.Button("Download", id='download_button')
                ]),
            dbc.Row(id='quiz-container'),
        ]),
    ])

    return layout


layout = serve_layout()


@callback(
    Output('input-container', 'children'),
    Input('input-now', 'n_submit'),
    State('input-container', 'children')
)
def update_input(n_submit, current_children):
    if n_submit is None:
        return current_children
    new_input = html.P(dcc.Input(
        id=f'input-now',
        type='text',
        value=''
    ))
    current_children.append(new_input)
    for i, child in enumerate(current_children[:-1]):
        child['props']['children']['props']['id'] = f'input-now-{i + 1}'
    return current_children


@callback(
    Output('read_status', 'children'),
    Input('submit_button', 'n_clicks'),
    State('input-container', 'children')
)
def submit_inputs(n_clicks, current_children):
    if n_clicks is None:
        return []
    global links
    list_links = [child['props']['children']['props']['value'] for child in current_children if child['props']['children']['props']['value']!='']
    links = list_links
    if len(list_links)==1 and (',' in list_links[0] or ' ' in list_links[0]):
        sep = ' ' if ' ' in list_links[0] else '\n' if '\n' in list_links[0] else ',' if ',' in list_links[0] else '.'
        more = list_links[0].split(sep)
        if len(more)>1:
            list_links = list_links[1:] + more
            links = list_links

    if len(list_links) == 0:
        list_links = [txt_file for txt_file in os.listdir(Config.wd_path) if '.txt' in txt_file]

    return [html.P("Links read: "+str(len(list_links)) + " : " + str(list_links))]


@callback(
    Output('process_status', 'children'),
    Input('read_status', 'children')
)
def process_links(read):
    if len(read) and len(read[0]['props']['children']):
        global links, lm
        util = Util()
        for link in links:
            util.save_link_content(link)
            lm = LanguageModel()
        return [html.P("Processing training data...")]
    return []

def check_if_trained(count):
    global links
    global num
    check = links == [x for x in os.listdir(Config.wd_path) if '.txt' in x]
    check = os.path.exists(Config.wd_path+'response.json')
    check = num == count
    return check

difficulties=()
@callback(
    Output('input-question-count', 'value'),
    [Input('low-question-count', 'value'),
     Input('mid-question-count', 'value'),
     Input('high-question-count', 'value')]
)
def total_count(l,m,h):
    if l=='' and m=='' and h=='':
        return
    low = int(l) if l!='' else 0
    mid = int(m) if m != '' else 0
    high = int(h) if h != '' else 0
    total = low + mid + high
    global difficulties
    difficulties=(l.strip(),m.strip(),h.strip())
    return str(total)

@callback(
    Output('quiz_status', 'children'),
    [Input('process_status', 'children'),
    Input('input-question-count', 'value')]
)
def generate_quiz(processing, count):
    if len(processing) and len(processing[0]['props']['children']):
        try:
            done = check_if_trained(count)
            response_json_path = Config.wd_path + 'response.json'
            quiz_json_path = Config.wd_path + 'quiz.json'
            global quiz
            global difficulties
            if not done:
                lm = LanguageModel()
                lm.train_llm(Config.wd_path)
                response = lm.query_with_prompt_file("assets/quiztemplate.json", count, difficulties)
                lm.save_response_as_json(response, response_json_path)
                quiz_json = lm.parse_quiz_response(response)
                lm.save_quiz_as_json(quiz_json, quiz_json_path)
                global num
                num = count
            else:
                with open(response_json_path, 'r') as response_file:
                    response_data = json.load(response_file)
                if not os.path.exists(quiz_json_path):
                    lm = LanguageModel()
                    quiz_json = lm.parse_quiz_response(response_data)
                    lm.save_quiz_as_json(quiz_json, quiz_json_path)
                else:
                    with open(quiz_json_path, 'r') as quiz_file:
                        quiz_json = json.load(quiz_file)

                quiz = quiz_json
            processing = []
            quiz = quiz_json
        except:
            print(str(sys.exc_info()))
            return [html.P("Quiz generation failure.")]
        return [html.P("Quiz generated.")]
    return []

callbacks = []

@callback(
    Output('quiz-container', 'children'),
    [Input('show_button', 'n_clicks'),
    Input('quiz_status', 'value')]
)
def render_quiz(n_submit, quiz_status):
    if n_submit is None or not quiz or quiz_status==[]:  # Check if no submission or quiz data is empty
        return None

    quiz_layout = html.Div([
        html.P(),
        html.H2(children=quiz['title']),
        *[html.Div([
            html.P(),
            html.Span(children=question['number']),html.Span(children=[" : "]),
            html.Span(children=question['text']),
            *[html.P(children=q) for q in question['options']],
            dcc.Input(id=f'answer-{question["number"]}', type='text', placeholder='Enter answer...', debounce=True),
            html.Div(id=f'correct-answer-{question["number"]}', children=[question['correct_answer']]),
            html.P(),
            html.P(),
            html.P()
        ]) for question in quiz['questions']]
    ])


    # if 'questions' in quiz:
    #     global callbacks
    #     for question in quiz['questions']:
    #         question_number = question['number']
    #
    #         @callback(
    #             Output(f'correct-answer-{question_number}', 'children'),
    #             [Input(f'answer-{question_number}', 'n_submit')],
    #             [State(f'answer-{question_number}', 'value')]
    #         )
    #         def check_answer(n_submit, user_answer, question_number=question_number):
    #             if n_submit:
    #                 correct_answer = next(
    #                     q['correct_answer']
    #                     for q in quiz['questions']
    #                     if q['number'] == question_number
    #                 )
    #
    #                 if user_answer.lower() == correct_answer.lower():
    #                     return f'Correct Answer: {correct_answer}'
    #                 else:
    #                     return f'Incorrect. Correct Answer: {correct_answer}'
    #             else:
    #                 return None
    #
    #         callbacks.append(check_answer)

    return quiz_layout


@callback(
    Output('download_link', 'data'),
    Input('download_button', 'n_clicks'),
    Input('quiz_name', 'value'),
)
def download_quiz(n_clicks, quiz_name):
    if n_clicks:
        if quiz_name:
            quiz_doc = Document()
            quiz_doc.add_heading(quiz['title'], 0)
            quiz_doc.add_paragraph('')
            quiz_doc.add_heading('Questions', 1)
            quiz_doc.add_paragraph('')
            for question in quiz['questions']:
                quiz_doc.add_paragraph('')
                quiz_doc.add_paragraph(f"{question['number']} {question['text']}")
                for option in question['options']:
                    quiz_doc.add_paragraph(option)
                quiz_doc.add_paragraph('')
                quiz_doc.add_paragraph('')
            quiz_doc.add_paragraph('')
            quiz_doc.add_heading('Answers', 1)
            for question in quiz['questions']:
                quiz_doc.add_paragraph(question['correct_answer'])
            quiz_filename = f"{quiz_name}.docx"
            quiz_doc.save(Config.collections+quiz_filename)
            # return dcc.send_file(quiz_filename, filename=quiz_filename)
            return None

    return None













@callback(
    Output('correct-answer', 'children'),
    [Input('answer', 'n_submit')],
    [State('answer', 'id')]
)
def check_answer(n_submit, input_id):
    if n_submit:
        question_number = input_id.split('-')[-1]
        user_answer = layout[input_id].value
        correct_answer = next(
            question['correct_answer']
            for question in quiz['questions']
            if question['number'] == question_number
        )

        if user_answer.lower() == correct_answer.lower():
            return f'Correct Answer: {correct_answer}'
        else:
            return f'Incorrect. Correct Answer: {correct_answer}'